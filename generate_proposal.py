"""
================================================================================
Project VERA - Project Proposal Generator
Generates the project proposal in DOCX format
================================================================================
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os


def create_proposal():
    """Generate the Project VERA proposal as a DOCX file."""
    doc = Document()
    
    # ---- Document Style Setup ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ---- Title Page ----
    for _ in range(6):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('PROJECT VERA')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Virtual Engineering Review Agent')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    
    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tagline.add_run(
        'A Multi-Agent System for Technical Document Auditing & Compliance — Any Industry'
    )
    run.font.size = Pt(14)
    run.italic = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Capstone Project Proposal\nData Science & AI Certification')
    run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('February 2026')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_page_break()
    
    # ---- Table of Contents ----
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        ('1.', 'Executive Summary'),
        ('2.', 'Problem Statement'),
        ('3.', 'Project Objectives'),
        ('4.', 'Proposed Solution'),
        ('5.', 'System Architecture'),
        ('6.', 'Agent Design'),
        ('7.', 'Security Implementation (RBAC)'),
        ('8.', 'Technology Stack'),
        ('9.', 'Data Model & Ingestion Pipeline'),
        ('10.', 'Implementation Timeline'),
        ('11.', 'Testing & Validation Strategy'),
        ('12.', 'Deliverables'),
        ('13.', 'Risk Assessment'),
        ('14.', 'Conclusion'),
    ]
    for num, item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(f'{num}  {item}')
        run.font.size = Pt(11)
    
    doc.add_page_break()
    
    # =====================================================================
    # 1. Executive Summary
    # =====================================================================
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'Project VERA v2.5 (Virtual Engineering Review Agent) is a production-grade, '
        'multi-agent AI system designed to automate technical document auditing, '
        'cross-reference checking, and compliance verification. It addresses the '
        'critical challenge of documentation drift in regulated industries by '
        'autonomously auditing live database records against official specifications '
        'and informal email communications.'
    )
    doc.add_paragraph(
        'The system features a "Surgical Router" architecture for domain-isolated processing, '
        '"Information Lock" grounding to eliminate hallucinations, and a deterministic '
        '"Triangulation Discrepancy Engine." VERA v2.5 supports local (Ollama) and '
        'cloud (Gemini, Groq) LLM backends, making it adaptable to diverse security '
        'and performance requirements.'
    )
    doc.add_paragraph(
        'VERA leverages state-of-the-art AI technologies — including LangGraph for '
        'multi-agent orchestration, ChromaDB for vector-based document storage, and '
        'Google Gemini for large language model capabilities — to create an intelligent '
        'document auditing system. The system ingests multiple document types '
        '(datasheets, emails, SOPs), applies Role-Based Access Control (RBAC), and '
        'automatically detects discrepancies between formal and informal documentation.'
    )
    doc.add_paragraph(
        'Key differentiators of VERA include:'
    )
    for item in [
        'Multi-agent architecture with specialized agents for different document types',
        'RBAC enforcement at the retrieval layer using metadata filtering',
        '🛡️ Information Lock protocol to ensure zero hallucinations',
        '🏝️ Immutable Domain Isolation to prevent unauthorized domain bleed',
        '🗄️ Enterprise-grade decoupling of structured (SQL) and unstructured (Vector) data',
        'Automatic discrepancy detection between datasheets and internal emails',
        'Human escalation pathway for unauthorized access attempts',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # =====================================================================
    # 2. Problem Statement
    # =====================================================================
    doc.add_heading('2. Problem Statement', level=1)
    doc.add_paragraph(
        'In many regulated industries, the gap between official documentation and '
        'actual engineering practices poses significant risks. Consider the following '
        'scenario from semiconductor manufacturing (the demo domain):'
    )
    
    # Add a highlighted scenario box
    scenario = doc.add_paragraph()
    scenario.paragraph_format.left_indent = Cm(1.5)
    run = scenario.add_run(
        'SCENARIO: The official RTX-9000 datasheet specifies a maximum voltage of 5.0V. '
        'However, an internal engineering email from the VP of Engineering states that '
        'the voltage must be immediately lowered to 3.3V due to accelerated electromigration '
        'in recent silicon lots. A junior engineer, unaware of this email, continues testing '
        'at 5.0V — potentially causing product failures, customer returns, and safety incidents.'
    )
    run.italic = True
    
    doc.add_paragraph(
        'This problem is compounded by several factors:'
    )
    for item in [
        'Informal decisions made via email are not captured in formal documentation systems',
        'Access to critical internal communications is not uniformly controlled',
        'No automated system exists to cross-reference datasheets against emails',
        'Junior engineers may not have visibility into restricted internal decisions',
        'Compliance audits cannot easily surface email-based process changes',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # =====================================================================
    # 3. Project Objectives
    # =====================================================================
    doc.add_heading('3. Project Objectives', level=1)
    doc.add_paragraph(
        'The primary objectives of Project VERA are:'
    )
    
    objectives = [
        ('O1: Document Integration', 
         'Build a unified document ingestion pipeline that processes datasheets, '
         'emails, and SOPs into a searchable vector store with rich metadata.'),
        ('O2: Intelligent Retrieval with RBAC',
         'Implement Role-Based Access Control at the retrieval layer, ensuring '
         'that document access is enforced based on user roles (Senior vs Junior).'),
        ('O3: Multi-Agent Orchestration',
         'Design a LangGraph-based multi-agent system with specialized agents for '
         'technical specifications, compliance/email review, and discrepancy detection.'),
        ('O4: Discrepancy Detection',
         'Automatically identify conflicts between formal documents (datasheets) and '
         'informal communications (emails), generating structured discrepancy reports.'),
        ('O5: Security & Escalation',
         'Implement a human escalation pathway that triggers when unauthorized access '
         'is detected or when the system has low confidence in its analysis.'),
    ]
    
    for title, desc in objectives:
        p = doc.add_paragraph()
        run = p.add_run(title + ': ')
        run.bold = True
        p.add_run(desc)
    
    # =====================================================================
    # 4. Proposed Solution
    # =====================================================================
    doc.add_heading('4. Proposed Solution', level=1)
    doc.add_paragraph(
        'VERA is designed as a Retrieval-Augmented Generation (RAG) system enhanced with '
        'multi-agent capabilities. The solution consists of two primary components:'
    )
    
    doc.add_heading('4.1 Data Ingestion Pipeline (ingestion.py)', level=2)
    doc.add_paragraph(
        'The ingestion pipeline handles document processing and storage:'
    )
    for item in [
        'Accepts multiple document types: Product Datasheets, Engineering Emails, SOPs',
        'Splits documents into semantically meaningful chunks using RecursiveCharacterTextSplitter',
        'Tags each chunk with rich metadata: source type, access level, department, document ID',
        'Generates vector embeddings using Google Gemini Embedding model',
        'Persists the indexed documents to ChromaDB for fast similarity search',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('4.2 Multi-Agent System (app.py)', level=2)
    doc.add_paragraph(
        'The core application implements a LangGraph state machine with specialized agents:'
    )
    for item in [
        'Surgical Router: Performs LLM-based NER and surgical routing to domain clusters',
        'DB Agent Cluster: Executes natural language to SQL queries on domain databases',
        'Official Docs Cluster: Extracts high-precision facts from datasheets and manuals',
        'Informal Docs Cluster: Researches engineering emails and informal memos',
        'Response Generator: Synthesizes comprehensive answers with source citations',
        'Triangulation Engine: Performs final audit and deterministic conflict detection',
        'Escalation Handler: Manages unauthorized access and out-of-domain queries',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    # =====================================================================
    # 5. System Architecture
    # =====================================================================
    doc.add_heading('5. System Architecture', level=1)
    doc.add_paragraph(
        'The VERA system follows a directed acyclic graph (DAG) architecture implemented '
        'using LangGraph\'s StateGraph. The workflow processes each query through the '
        'following stages:'
    )
    
    # Architecture flow as a table
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['Stage', 'Component', 'Description']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    flow_data = [
        ('1', 'Entry Point', 'User submits query with role (senior/junior)'),
        ('2', 'Surgical Router', 'LLM-NER (Entity/Attr) + Cluster Delegation'),
        ('3a', 'DB Cluster', 'SQL Querying (Read-only) if db_query target'),
        ('3b', 'Official Cluster', 'Spec Extraction if spec_retrieval target'),
        ('3c', 'Full Retrieval', 'Cross-referencing all sources if cross_reference'),
        ('4', 'Triangulation', 'Deterministic audit using authority hierarchy'),
    ]
    for i, (stage, component, desc) in enumerate(flow_data):
        table.rows[i+1].cells[0].text = stage
        table.rows[i+1].cells[1].text = component
        table.rows[i+1].cells[2].text = desc
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The conditional routing mechanism ensures that queries are directed to the '
        'appropriate agent based on intent classification, while the RBAC layer filters '
        'documents before they reach the generation stage.'
    )
    
    # =====================================================================
    # 6. Agent Design
    # =====================================================================
    doc.add_heading('6. Agent Design', level=1)
    doc.add_paragraph(
        'Each agent in VERA is implemented as a node in the LangGraph StateGraph. '
        'The agents share a common state schema (GraphState) and communicate through '
        'state updates.'
    )
    
    # Agent details table
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    agent_headers = ['Agent', 'Node Name', 'Input', 'Output']
    for i, header in enumerate(agent_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    agents_data = [
        ('Surgical Router', 'route_query', 'Question + Role', 'Route + NER + Cluster Flag'),
        ('DB Cluster', '{domain}_db_query', 'Question', 'SQL Results (Priority 3)'),
        ('Official Docs', '{domain}_official', 'Question + Entity', 'Specs (Priority 2)'),
        ('Informal Docs', '{domain}_informal', 'Question + Entity', 'Memos (Priority 1)'),
        ('Generator', 'generate_response', 'Extracted Facts', 'LLM Report (Grounded)'),
        ('Triangulation', '{domain}_discrepancy', 'All Facts', 'Verdict (Deterministic)'),
    ]
    for i, (agent, node, inp, out) in enumerate(agents_data):
        table.rows[i+1].cells[0].text = agent
        table.rows[i+1].cells[1].text = node
        table.rows[i+1].cells[2].text = inp
        table.rows[i+1].cells[3].text = out
    
    doc.add_paragraph()
    
    doc.add_heading('6.1 GraphState Schema', level=2)
    doc.add_paragraph(
        'The shared state is defined as a TypedDict with the following fields:'
    )
    state_items = [
        'question (str): The user\'s input query',
        'generation (str): The LLM-generated response',
        'user_role (str): "senior" or "junior" — determines access level',
        'user_domain (str): The user\'s assigned domain (e.g., "semiconductor", "medical")',
        'documents (List[Document]): Retrieved documents from ChromaDB',
        'route (str): Routing decision — "technical", "compliance", or "escalate"',
        'flagged (bool): Security flag for unauthorized access attempts',
        'metadata_log (str): Audit log of retrieval operations',
        'retrieved_docs (dict): Per-agent docs for cross-agent comparison (tech vs compliance)',
        'discrepancy_report (str): Structured report from the Case Agent',
        'next_agent (str): Domain-based routing target (e.g., "semiconductor")',
        'refinement_count (int): Tracks discussion loop iterations',
        'max_refinements (int): Configurable limit for discussion loop',
        'critique (str): Feedback from Discrepancy Agent to Response Agent',
    ]
    for item in state_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # =====================================================================
    # 7. Security Implementation (RBAC)
    # =====================================================================
    doc.add_heading('7. Security Implementation (RBAC)', level=1)
    doc.add_paragraph(
        'VERA implements a three-layer security system to enforce Role-Based Access Control:'
    )
    
    doc.add_heading('7.1 Layer 1 — Hybrid Router Security', level=2)
    doc.add_paragraph(
        'The Router Agent employs a hybrid classification strategy: deterministic '
        'keyword matching for speed, followed by LLM-based intent fallback for '
        'ambiguous queries. For security, the LLM analyzes whether a junior user\'s '
        'query implies access to restricted information. If flagged, the query is '
        'routed directly to the escalation handler.'
    )
    
    doc.add_heading('7.2 Layer 2 — Metadata Filtering', level=2)
    doc.add_paragraph(
        'This is the core RBAC mechanism. Every document in ChromaDB is tagged with an '
        'access_level metadata field (public, internal_only, or confidential). When a '
        'retrieval query is executed, the system dynamically constructs a metadata filter '
        'based on the user\'s role:'
    )
    
    # RBAC access matrix
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    
    rbac_headers = ['Access Level', 'Senior (User A)', 'Junior (User B)']
    for i, header in enumerate(rbac_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    rbac_data = [
        ('public', '✓ Full Access', '✓ Full Access'),
        ('internal_only', '✓ Full Access', '✗ Blocked'),
        ('confidential', '✓ Full Access', '✗ Blocked + Escalated'),
    ]
    for i, (level, senior, junior) in enumerate(rbac_data):
        table.rows[i+1].cells[0].text = level
        table.rows[i+1].cells[1].text = senior
        table.rows[i+1].cells[2].text = junior
    
    doc.add_paragraph()
    
    doc.add_heading('7.3 Layer 3 — Domain Isolation', level=2)
    doc.add_paragraph(
        'To prevent "Domain Bleed," VERA treats the user\'s assigned domain as '
        'immutable. Even if a query contains keywords from another domain, the system '
        'strictly restricts retrieval to the authorized domain, escalating any detection '
        'of cross-domain intent to a supervisor node.'
    )

    doc.add_heading('7.4 Layer 4 — Information Lock (Grounding)', level=2)
    doc.add_paragraph(
        'The "Information Lock" is a grounding-first protocol that ensures the system '
        'only responds using facts explicitly found in provided documents or database '
        'records. If the required information is missing, the system is hard-coded to '
        'respond with a standardized "Data Not Found" message rather than inferring '
        'or using external training data.'
    )

    doc.add_heading('7.5 Layer 5 — Escalation', level=2)
    doc.add_paragraph(
        'When security checks or domain isolation flags a query, the workflow bypasses '
        'retrieval and generation, routing directly to the escalation handler.'
    )
    
    # =====================================================================
    # 8. Technology Stack
    # =====================================================================
    doc.add_heading('8. Technology Stack', level=1)
    
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Light Grid Accent 1'
    
    tech_headers = ['Technology', 'Purpose', 'Version']
    for i, header in enumerate(tech_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    tech_data = [
        ('LangGraph', 'Multi-agent state machine orchestration', '≥ 0.2.0'),
        ('LangChain', 'RAG pipeline, prompt management', '≥ 0.3.0'),
        ('ChromaDB', 'Local vector store with metadata filtering', '≥ 0.5.0'),
        ('Streamlit', 'Interactive chat UI with agent feedback', '≥ 1.40.0'),
        ('Google Gemini', 'Cloud LLM + Embeddings (Option A)', 'Latest'),
        ('Groq', 'High-speed cloud inference (Option B)', 'Latest'),
        ('Ollama', 'Local LLM + Embeddings (Option C)', 'Latest'),
        ('Python', 'Core runtime', '3.10+'),
    ]
    for i, (tech, purpose, version) in enumerate(tech_data):
        table.rows[i+1].cells[0].text = tech
        table.rows[i+1].cells[1].text = purpose
        table.rows[i+1].cells[2].text = version
    
    doc.add_paragraph()
    doc.add_paragraph(
        'VERA supports a multi-backend architecture: users can choose between Google Gemini '
        '(cloud API), Groq (high-speed Llama-3.3 inference), or Ollama (100% local execution). '
        'The LLM_BACKEND environment variable controls which backend is used for inference.'
    )
    
    # =====================================================================
    # 9. Data Model & Ingestion Pipeline
    # =====================================================================
    doc.add_heading('9. Data Model & Ingestion Pipeline', level=1)
    
    doc.add_heading('9.1 Document Types', level=2)
    doc.add_paragraph(
        'VERA processes three categories of documents, each with distinct metadata:'
    )
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Light Grid Accent 1'
    
    doc_headers = ['Document Type', 'Source Tag', 'Typical Access', 'Example']
    for i, header in enumerate(doc_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc_data = [
        ('Product Datasheet', 'datasheet', 'public', 'RTX-9000 voltage specs'),
        ('Engineering Email', 'email', 'internal/confidential', 'Voltage limit change decision'),
        ('SOP', 'sop', 'public/internal', 'Wafer testing procedure'),
        ('DB Info', 'db_info', 'public/confidential', 'Production lot data, test results'),
        ('Versioned Document', 'document', 'public/confidential', 'Specification revisions (v1/v2)'),
    ]
    for i, (dtype, source, access, example) in enumerate(doc_data):
        table.rows[i+1].cells[0].text = dtype
        table.rows[i+1].cells[1].text = source
        table.rows[i+1].cells[2].text = access
        table.rows[i+1].cells[3].text = example
    
    doc.add_paragraph()
    
    doc.add_heading('9.2 Ingestion Pipeline Steps', level=2)
    steps = [
        'Document Creation: Mock data simulates 8 semiconductor documents',
        'Text Splitting: RecursiveCharacterTextSplitter (500 char chunks, 50 overlap)',
        'Metadata Tagging: Each chunk inherits source, access_level, department tags',
        'Embedding Generation: Google Gemini embedding model converts text to vectors',
        'Vector Store Persistence: ChromaDB stores vectors + metadata locally',
    ]
    for i, step in enumerate(steps):
        p = doc.add_paragraph()
        run = p.add_run(f'Step {i+1}: ')
        run.bold = True
        p.add_run(step)
    
    # =====================================================================
    # 10. Implementation Timeline
    # =====================================================================
    doc.add_heading('10. Implementation Timeline', level=1)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Light Grid Accent 1'
    
    timeline_headers = ['Phase', 'Duration', 'Deliverables']
    for i, header in enumerate(timeline_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    timeline_data = [
        ('Phase 1: Research & Design', '1 week', 'Architecture diagram, technology selection'),
        ('Phase 2: Data Ingestion', '1 week', 'ingestion.py, ChromaDB setup, mock data'),
        ('Phase 3: Agent Development', '2 weeks', 'app.py, LangGraph workflow, RBAC, agents'),
        ('Phase 4: Testing & Validation', '1 week', 'Test scenarios, discrepancy detection'),
        ('Phase 5: Documentation', '3 days', 'README.md, project proposal, walkthrough'),
    ]
    for i, (phase, duration, deliverables) in enumerate(timeline_data):
        table.rows[i+1].cells[0].text = phase
        table.rows[i+1].cells[1].text = duration
        table.rows[i+1].cells[2].text = deliverables
    
    # =====================================================================
    # 11. Testing & Validation Strategy
    # =====================================================================
    doc.add_heading('11. Testing & Validation Strategy', level=1)
    doc.add_paragraph(
        'VERA is validated through four primary test scenarios that demonstrate '
        'the system\'s core capabilities:'
    )
    
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    
    test_headers = ['Test #', 'User Role', 'Query', 'Expected Behavior']
    for i, header in enumerate(test_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    test_data = [
        ('1', 'Senior', 'RTX-9000 voltage limit', 'Full info including internal email (3.3V)'),
        ('2', 'Junior', 'RTX-9000 voltage limit', 'Public datasheet only (5.0V)'),
        ('3', 'Junior', 'Internal burn-in emails', 'ESCALATED — access denied'),
        ('4', 'Senior', 'Quality audit + email changes', 'SOPs + emails with discrepancy report'),
        ('5', 'Senior', 'Spec versions + production DB', 'DB records + versioned docs, version discrepancy'),
        ('6', 'Senior (semi)', 'FDA clinical trial requirements', 'ESCALATED — out-of-domain query'),
    ]
    for i, (num, role, query, expected) in enumerate(test_data):
        table.rows[i+1].cells[0].text = num
        table.rows[i+1].cells[1].text = role
        table.rows[i+1].cells[2].text = query
        table.rows[i+1].cells[3].text = expected
    
    # =====================================================================
    # 12. Deliverables
    # =====================================================================
    doc.add_heading('12. Deliverables', level=1)
    
    deliverables = [
        ('streamlit_app.py', 'Interactive Streamlit chat interface with role switching, '
         'document viewer, RBAC audit log, discrepancy reports, and agent execution trace'),
        ('ingestion.py', 'Data ingestion pipeline with mock data, text splitting, '
         'embedding generation, and ChromaDB persistence with RBAC metadata tagging'),
        ('app.py', 'Main LangGraph multi-agent application with Router, Tech Spec Agent, '
         'Compliance Agent, Generator, Case Agent (discrepancy detector), and Escalation handler'),
        ('README.md', 'Professional documentation with architecture diagram (Mermaid.js), '
         'agent descriptions, RBAC security explanation, and setup instructions'),
        ('environment.yml', 'Conda environment definition for reproducible setup'),
        ('requirements.txt', 'Python dependencies (alternative to conda)'),
        ('TEAM_GUIDE.md', 'Developer guide for team members with templates and conventions'),
        ('Project Proposal (DOCX)', 'Comprehensive project proposal document'),
        ('.env', 'Environment configuration for backend and API keys'),
    ]
    for filename, desc in deliverables:
        p = doc.add_paragraph()
        run = p.add_run(filename + ': ')
        run.bold = True
        p.add_run(desc)
    
    # =====================================================================
    # 13. Risk Assessment
    # =====================================================================
    doc.add_heading('13. Risk Assessment', level=1)
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Light Grid Accent 1'
    
    risk_headers = ['Risk', 'Likelihood', 'Impact', 'Mitigation']
    for i, header in enumerate(risk_headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    risks = [
        ('API Rate Limiting', 'High', 'Medium', 
         'Exponential backoff retry logic implemented'),
        ('Embedding Model Deprecation', 'Medium', 'High',
         'Modular design allows quick model swaps'),
        ('RBAC Bypass via Prompt Injection', 'Low', 'High',
         'Metadata filtering at DB layer, not prompt-level'),
        ('LLM Hallucination in Routing', 'Low', 'Low',
         'Router uses deterministic keyword matching, not LLM'),
    ]
    for i, (risk, likelihood, impact, mitigation) in enumerate(risks):
        table.rows[i+1].cells[0].text = risk
        table.rows[i+1].cells[1].text = likelihood
        table.rows[i+1].cells[2].text = impact
        table.rows[i+1].cells[3].text = mitigation
    
    # =====================================================================
    # 14. Conclusion
    # =====================================================================
    doc.add_heading('14. Conclusion', level=1)
    doc.add_paragraph(
        'Project VERA demonstrates a production-grade approach to solving the critical '
        'problem of document inconsistency in any technical or regulated industry. '
        'While the included demo highlights a semiconductor use case, the architecture '
        'is fully domain-agnostic. By combining LangGraph multi-agent orchestration, '
        'ChromaDB vector storage with metadata filtering, and Google Gemini\'s LLM '
        'capabilities, VERA provides:'
    )
    for item in [
        'Intelligent document retrieval with Role-Based Access Control',
        'Automatic detection of discrepancies between formal and informal documentation',
        'A security-first architecture with human escalation for unauthorized access',
        'Email context analysis to surface informal engineering decisions',
        'A modular, extensible framework that can be adapted to any document-heavy industry',
    ]:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph(
        'The system is designed to be immediately applicable to any document-heavy industry, '
        'with the mock data easily replaceable by actual document loaders for '
        'PDFs, email archives, and SOP management systems. VERA represents a significant '
        'step forward in applying AI-powered multi-agent systems to industrial document '
        'compliance and auditing across sectors such as semiconductor manufacturing, '
        'aerospace, pharmaceuticals, automotive, energy, and financial services.'
    )
    
    # ---- Save the document ----
    output_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        'Project_VERA_Proposal.docx'
    )
    doc.save(output_path)
    print(f"[SUCCESS] Project proposal saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_proposal()
